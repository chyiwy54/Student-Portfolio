import os, base64, requests
from flask import Flask, render_template, request, send_file, redirect, url_for, flash
from io import BytesIO
from docx import Document
from docx.shared import Inches

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "dev-secret")  # for flash

# 1) 讀取 API Key（不要硬寫在檔案）
GEMINI_API_KEY = "AIzaSyDEjjmrDEQFLxdUA0OsiCGVwUVvgVr7r8w"
if not GEMINI_API_KEY:
    print("[警告] 尚未設定 GEMINI_API_KEY 環境變數，Gemini 功能將無法使用。")
GEMINI_ENDPOINT = f"https://generativelanguage.googleapis.com/v1/models/gemini-1.5-flash:generateContent?key={GEMINI_API_KEY}"


# 2) 將前端表單內容整理成一段清楚的文字，交給 Gemini
def build_user_prompt(form):
    lines = []
    def add(label, key):
        v = (form.get(key) or "").strip()
        if v:
            lines.append(f"{label}：{v}")
    add("作品主題", "topic")
    add("作者/班級", "author")
    add("摘要", "abstract")
    add("主題說明", "topic_desc")
    add("成果展示", "result")
    add("心得反思", "reflection")
    add("參考資料", "references")
    return "\n".join(lines).strip()


# 3) 呼叫 Gemini（支援多張圖片）
def call_gemini_api(images_bytes_list, user_input):
    image_parts = [
        {"inlineData": {"mimeType": "image/jpeg", "data": base64.b64encode(b).decode("utf-8")}}
        for b in images_bytes_list if b
    ]

    system_prompt = (
        "請使用繁體中文，根據以下活動資訊，產生有標題的小節，至少包含：\n"
        "一、主題說明\n二、心得反思\n三、學習歷程檔案內容簡介（100字內）\n"
        "四、檢討或反思\n五、學習或執行過程（步驟、結果）\n六、對未來的影響\n七、學習成果佐證說明\n"
        "語氣自然、條列清楚，可用粗體小標。"
    )

    payload = {
        "contents": [{
            "parts": [
                *image_parts,
                {"text": system_prompt + "\n\n活動資訊如下：\n" + (user_input or "（使用者未提供補充說明）")}
            ]
        }]
    }
    headers = {"Content-Type": "application/json"}

    resp = requests.post(GEMINI_ENDPOINT, json=payload, headers=headers)
    try:
        data = resp.json()
    except Exception:
        return f"[錯誤] 呼叫 Gemini 失敗，HTTP {resp.status_code}"

    if "candidates" in data and data["candidates"]:
        return data["candidates"][0]["content"]["parts"][0]["text"]
    return "[錯誤] Gemini 回傳內容如下：\n" + str(data)


@app.route("/", methods=["GET"])
def index():
    return render_template("index.html")


@app.route("/generate", methods=["POST"])
def generate():
    if not GEMINI_API_KEY:
        flash("尚未設定 GEMINI_API_KEY，請先在環境變數中設定。")
        return redirect(url_for("index"))

    # 1) 取表單欄位
    images = request.files.getlist("images")  # 對應前端 name="images"
    output = request.form.get("output", "docx")  # pdf or docx
    want_preview = request.form.get("preview") == "1"

    # 2) 整理成 Gemini 用的文字，避免過少
    user_prompt = build_user_prompt(request.form)
    if len(user_prompt) < 10:  # 無硬性字數限制，但太少就提醒
        flash("內容太少了，請至少補充 1～2 句，讓系統知道你要什麼。")
        return redirect(url_for("index"))

    # 3) 讀圖片 bytes（避免 payload 過大，先取最多 3 張；要全部就改成 imgs_bytes）
    imgs_bytes = []
    for f in images:
        if f and getattr(f, "filename", ""):
            imgs_bytes.append(f.read())
    imgs_bytes = imgs_bytes[:3]

    # 4) 呼叫 Gemini 生成正文
    generated = call_gemini_api(imgs_bytes, user_prompt)

    # 5a) 若是「預覽」：直接顯示 Gemini 草稿（不產檔）
    if want_preview:
        return render_template("preview.html", content=generated)

    # 5b) 產 DOCX；圖片統一放最後
    doc = Document()
    doc.add_heading("學習歷程紀錄", 0)

    # 你原本是用 ** 切段；若模型輸出含 Markdown，可直接整段貼上
    for block in generated.split("**"):
        text = block.strip()
        if text:
            doc.add_paragraph(text)

    # 分頁＋圖片
    doc.add_page_break()
    doc.add_heading("活動照片", level=1)
    for f in images:
        if f and getattr(f, "filename", ""):
            try:
                f.stream.seek(0)
                doc.add_picture(f, width=Inches(4))
                doc.add_paragraph("")
            except Exception:
                # 不支援的格式可略過或之後用 Pillow 轉成 JPEG
                pass

    # 6) 存到記憶體緩衝並回傳
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    if output == "docx":
        return send_file(
            buf,
            as_attachment=True,
            download_name="learning.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # PDF：先回退為 DOCX（等你安裝 LibreOffice 或改走 HTML→PDF 再開）
    return send_file(
        buf,
        as_attachment=True,
        download_name="learning.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    # 可改成從環境讀取 PORT，方便雲端部署
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", 5000)), debug=True)
